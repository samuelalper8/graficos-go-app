"""
Módulo de exportação para múltiplos formatos
"""

import io
import os
from datetime import datetime
from pathlib import Path
import matplotlib.pyplot as plt
from matplotlib.figure import Figure
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Image, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import zipfile


class ExportadorWord:
    """Exporta gráficos para Word (DOCX)"""
    
    def __init__(self, titulo="Relatório de Gráficos"):
        self.doc = Document()
        self.titulo = titulo
        self._setup_documento()
    
    def _setup_documento(self):
        """Configura estilos e cabeçalho do documento"""
        # Adicionar título
        titulo = self.doc.add_heading(self.titulo, 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar data
        data_paragrafo = self.doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        data_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Espaço
    
    def adicionar_grafico(self, figura: Figure, municipio: str, serie: str):
        """Adiciona um gráfico ao documento"""
        # Título da seção
        heading = self.doc.add_heading(f"{municipio} - {serie}", level=2)
        
        # Salvar figura em bytes
        img_bytes = io.BytesIO()
        figura.savefig(img_bytes, format='png', dpi=150, bbox_inches='tight')
        img_bytes.seek(0)
        
        # Adicionar imagem ao documento
        self.doc.add_picture(img_bytes, width=Inches(6))
        
        # Adicionar quebra de página
        self.doc.add_page_break()
    
    def salvar(self, caminho: str):
        """Salva o documento"""
        self.doc.save(caminho)
        return caminho


class ExportadorPDF:
    """Exporta gráficos para PDF"""
    
    def __init__(self, titulo="Relatório de Gráficos"):
        self.titulo = titulo
        self.figuras = []
    
    def adicionar_grafico(self, figura: Figure, municipio: str, serie: str):
        """Adiciona um gráfico à lista"""
        self.figuras.append({
            'figura': figura,
            'municipio': municipio,
            'serie': serie
        })
    
    def salvar(self, caminho: str):
        """Salva o PDF"""
        from reportlab.platypus import SimpleDocTemplate, Image, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import letter
        
        # Criar documento
        doc = SimpleDocTemplate(caminho, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#0072B2',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph(self.titulo, title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Data
        data_style = ParagraphStyle(
            'Data',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER
        )
        story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", data_style))
        story.append(Spacer(1, 0.3*inch))
        story.append(PageBreak())
        
        # Adicionar gráficos
        for item in self.figuras:
            # Título do gráfico
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12
            )
            story.append(Paragraph(f"{item['municipio']} - {item['serie']}", heading_style))
            
            # Salvar figura em bytes
            img_bytes = io.BytesIO()
            item['figura'].savefig(img_bytes, format='png', dpi=150, bbox_inches='tight')
            img_bytes.seek(0)
            
            # Adicionar imagem
            img = Image(img_bytes, width=6*inch, height=4*inch)
            story.append(img)
            story.append(Spacer(1, 0.3*inch))
            story.append(PageBreak())
        
        # Construir PDF
        doc.build(story)
        return caminho


class ExportadorPNG:
    """Exporta gráficos para PNG em ZIP"""
    
    def __init__(self):
        self.figuras = []
    
    def adicionar_grafico(self, figura: Figure, municipio: str, serie: str):
        """Adiciona um gráfico à lista"""
        self.figuras.append({
            'figura': figura,
            'municipio': municipio,
            'serie': serie
        })
    
    def salvar(self, caminho_zip: str):
        """Salva todos os gráficos em um ZIP"""
        with zipfile.ZipFile(caminho_zip, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for idx, item in enumerate(self.figuras):
                # Nome do arquivo
                nome_arquivo = f"{item['municipio']}_{item['serie'].replace(' ', '_')}.png"
                
                # Salvar figura em bytes
                img_bytes = io.BytesIO()
                item['figura'].savefig(img_bytes, format='png', dpi=150, bbox_inches='tight')
                img_bytes.seek(0)
                
                # Adicionar ao ZIP
                zipf.writestr(nome_arquivo, img_bytes.getvalue())
        
        return caminho_zip


def gerar_todos_formatos(graficos_r2000, graficos_r4000, diretorio_saida=None):
    """
    Gera arquivos em todos os formatos
    
    Args:
        graficos_r2000: Dict com gráficos R-2000
        graficos_r4000: Dict com gráficos R-4000
        diretorio_saida: Diretório para salvar os arquivos
    
    Returns:
        Dict com caminhos dos arquivos gerados
    """
    if diretorio_saida is None:
        diretorio_saida = Path(tempfile.gettempdir()) / "graficos_go"
        diretorio_saida.mkdir(exist_ok=True)
    
    arquivos = {}
    
    # Exportar Word
    exportador_word = ExportadorWord("Relatório de Monitoramento GO")
    if graficos_r2000:
        for municipio, fig in graficos_r2000.items():
            exportador_word.adicionar_grafico(fig, municipio, "R-2000")
    if graficos_r4000:
        for municipio, fig in graficos_r4000.items():
            exportador_word.adicionar_grafico(fig, municipio, "R-4000")
    
    caminho_word = os.path.join(diretorio_saida, "Relatorio_Graficos.docx")
    exportador_word.salvar(caminho_word)
    arquivos['word'] = caminho_word
    
    # Exportar PDF
    exportador_pdf = ExportadorPDF("Relatório de Monitoramento GO")
    if graficos_r2000:
        for municipio, fig in graficos_r2000.items():
            exportador_pdf.adicionar_grafico(fig, municipio, "R-2000")
    if graficos_r4000:
        for municipio, fig in graficos_r4000.items():
            exportador_pdf.adicionar_grafico(fig, municipio, "R-4000")
    
    caminho_pdf = os.path.join(diretorio_saida, "Relatorio_Graficos.pdf")
    exportador_pdf.salvar(caminho_pdf)
    arquivos['pdf'] = caminho_pdf
    
    # Exportar PNGs em ZIP
    exportador_png = ExportadorPNG()
    if graficos_r2000:
        for municipio, fig in graficos_r2000.items():
            exportador_png.adicionar_grafico(fig, municipio, "R-2000")
    if graficos_r4000:
        for municipio, fig in graficos_r4000.items():
            exportador_png.adicionar_grafico(fig, municipio, "R-4000")
    
    caminho_zip = os.path.join(diretorio_saida, "Graficos_PNG.zip")
    exportador_png.salvar(caminho_zip)
    arquivos['png_zip'] = caminho_zip
    
    return arquivos
